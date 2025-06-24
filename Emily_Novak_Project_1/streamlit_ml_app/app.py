# STREAMLIT ML CLASSIFICATION APP - TRIPLE MODEL SUPPORT
# =====================================================

import streamlit as st
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
import seaborn as sns
import pdfplumber
import pathlib
from docx import Document

# Page Configuration
st.set_page_config(
    page_title="ML Text Classifier",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #007bff;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# MODEL LOADING SECTION
# ============================================================================

@st.cache_resource
def load_models():
    models = {}
    
    try:
        # Load the main pipeline (Logistic Regression)
        try:
            models['pipeline'] = joblib.load('models/ai_detection_pipeline.pkl')
            models['pipeline_available'] = True
        except FileNotFoundError:
            models['pipeline_available'] = False
        
        # Load TF-IDF vectorizer
        try:
            models['vectorizer'] = joblib.load('models/tfidf_vectorizer.pkl')
            models['vectorizer_available'] = True
        except FileNotFoundError:
            models['vectorizer_available'] = False
        
        # Load SVM model
        try:
            models['svm'] = joblib.load('models/svm_model.pkl')
            models['svm_available'] = True
        except FileNotFoundError:
            models['svm_available'] = False
        
        # Load Decision Tree model
        try:
            models['decision_tree'] = joblib.load('models/dt_model.pkl')
            models['dt_available'] = True
        except FileNotFoundError:
            models['dt_available'] = False

        # Load AdaBoost model
        try:
            models['adaboost'] = joblib.load('models/ada_model.pkl')
            models['ada_available'] = True
        except FileNotFoundError:
            models['ada_available'] = False
        
        # Check if at least one complete setup is available
        pipeline_ready = models['pipeline_available']
        individual_ready = models['vectorizer_available'] and (models['svm_available'] or models['dt_available'] or models['ada_available'])
        
        if not (pipeline_ready or individual_ready):
            st.error("No complete model setup found!")
            return None
        
        return models
        
    except Exception as e:
        st.error(f"Error loading models: {e}")
        return None

# ============================================================================
# PREDICTION FUNCTION
# ============================================================================

def make_prediction(text, model_choice, models):
    """Make prediction using the selected model"""
    if models is None:
        return None, None
    
    try:
        prediction = None
        probabilities = None
        
        if model_choice == "pipeline" and models.get('pipeline_available'):
            # Use the complete pipeline (Logistic Regression)
            prediction = models['pipeline'].predict([text])[0]
            probabilities = models['pipeline'].predict_proba([text])[0]
            
        elif model_choice == "svm":
            if models.get('pipeline_available'):
                # Use pipeline for LR
                prediction = models['pipeline'].predict([text])[0]
                probabilities = models['pipeline'].predict_proba([text])[0]
            elif models.get('vectorizer_available') and models.get('svm_available'):
                # Use individual components
                X = models['vectorizer'].transform([text])
                prediction = models['svm'].predict(X)[0]
                probabilities = models['svm'].predict_proba(X)[0]
                
        elif model_choice == "decision_tree":
            if models.get('vectorizer_available') and models.get('dt_available'):
                # Use individual components for NB
                X = models['vectorizer'].transform([text])
                prediction = models['decision_tree'].predict(X)[0]
                probabilities = models['decision_tree'].predict_proba(X)[0]
        elif model_choice == "adaboost":
            if models.get('vectorizer_available') and models.get('ada_available'):
                # Use individual components for NB
                X = models['vectorizer'].transform([text])
                prediction = models['adaboost'].predict(X)[0]
                probabilities = models['adaboost'].predict_proba(X)[0]
        
        if prediction is not None and probabilities is not None:
            # Convert to readable format
            class_names = ['Human', 'AI']
            prediction_label = class_names[prediction]
            return prediction_label, probabilities
        else:
            return None, None
        
    except Exception as e:
        st.error(f"Error making prediction: {e}")
        st.error(f"Model choice: {model_choice}")
        st.error(f"Available models: {[k for k, v in models.items() if isinstance(v, bool) and v]}")
        return None, None

def get_available_models(models):
    """Get list of available models for selection"""
    available = []
    
    if models is None:
        return available
    
    if models.get('pipeline_available'):
        available.append(("svm", "📈 SVM (Pipeline)"))
    elif models.get('vectorizer_available') and models.get('svm_available'):
        available.append(("svm", "📈 SVM (Individual)"))
    
    if models.get('vectorizer_available') and models.get('dt_available'):
        available.append(("decision_tree", "🎯 Decision Tree"))

    if models.get('vectorizer_available') and models.get('ada_available'):
        available.append(("adaboost", "🎯 AdaBoost"))
    
    return available

# ============================================================================
# SIDEBAR NAVIGATION
# ============================================================================

st.sidebar.title("🧭 Navigation")
st.sidebar.markdown("Choose what you want to do:")

page = st.sidebar.selectbox(
    "Select Page:",
    ["🏠 Home", "🔮 Single Prediction", "📁 Batch Processing", "⚖️ Model Comparison", "📊 Model Info", "❓ Help"]
)

# Load models
models = load_models()

# ============================================================================
# HOME PAGE
# ============================================================================

if page == "🏠 Home":
    st.markdown('<h1 class="main-header">🤖 ML Text Classification App</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    Welcome to your machine learning web application! This app demonstrates AI detection
    using multiple trained models: **SVM**, **Decision Tree**, and **AdaBoost**.
    """)
    
    # App overview
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        ### 🔮 Single Prediction
        - Enter text manually
        - Choose between models
        - Get instant predictions
        - See confidence scores
        """)
    
    with col2:
        st.markdown("""
        ### 📁 Batch Processing
        - Upload text files
        - Process multiple texts
        - Compare model performance
        - Download results
        """)
    
    with col3:
        st.markdown("""
        ### ⚖️ Model Comparison
        - Compare different models
        - Side-by-side results
        - Agreement analysis
        - Performance metrics
        """)
    
    # Model status
    st.subheader("📋 Model Status")
    if models:
        st.success("✅ Models loaded successfully!")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if models.get('pipeline_available'):
                st.info("**📈 SVM**\n✅ Pipeline Available")
            elif models.get('svm_available') and models.get('vectorizer_available'):
                st.info("**📈 SVM**\n✅ Individual Components")
            else:
                st.warning("**📈 Logistic Regression**\n❌ Not Available")
        
        with col2:
            if models.get('dt_available') and models.get('vectorizer_available'):
                st.info("**🎯 Decision Tree**\n✅ Available")
            else:
                st.warning("**🎯 Decision Tree**\n❌ Not Available")

        with col3:
            if models.get('ada_available') and models.get('vectorizer_available'):
                st.info("**🎯 AdaBoost**\n✅ Available")
            else:
                st.warning("**🎯 AdaBoost**\n❌ Not Available")
        
        with col4:
            if models.get('vectorizer_available'):
                st.info("**🔤 TF-IDF Vectorizer**\n✅ Available")
            else:
                st.warning("**🔤 TF-IDF Vectorizer**\n❌ Not Available")
        
    else:
        st.error("❌ Models not loaded. Please check model files.")

# ============================================================================
# SINGLE PREDICTION PAGE
# ============================================================================

elif page == "🔮 Single Prediction":
    st.header("🔮 Make a Single Prediction")
    st.markdown("Enter text below and select a model to detect AI-written text.")
    
    if models:
        available_models = get_available_models(models)
        
        if available_models:
            # Model selection
            model_choice = st.selectbox(
                "Choose a model:",
                options=[model[0] for model in available_models],
                format_func=lambda x: next(model[1] for model in available_models if model[0] == x)
            )
            
            # Text input
            user_input = st.text_area(
                "Enter your text here:",
                placeholder="Type or paste your text here (e.g., product review, feedback, comment)...",
                height=150
            )
            
            # Character count
            if user_input:
                st.caption(f"Character count: {len(user_input)} | Word count: {len(user_input.split())}")
            
            # Example texts
            with st.expander("📝 Try these example texts"):
                examples = [
                    "This product is absolutely amazing! Best purchase I've made this year.",
                    "Terrible quality, broke after one day. Complete waste of money.",
                    "It's okay, nothing special but does the job.",
                    "Outstanding customer service and fast delivery. Highly recommend!",
                    "I love this movie! It's absolutely fantastic and entertaining."
                ]
                
                col1, col2 = st.columns(2)
                for i, example in enumerate(examples):
                    with col1 if i % 2 == 0 else col2:
                        if st.button(f"Example {i+1}", key=f"example_{i}"):
                            st.session_state.user_input = example
                            st.rerun()
            
            # Use session state for user input
            if 'user_input' in st.session_state:
                user_input = st.session_state.user_input
            
            # Prediction button
            if st.button("🚀 Predict", type="primary"):
                if user_input.strip():
                    with st.spinner('Analyzing text...'):
                        prediction, probabilities = make_prediction(user_input, model_choice, models)
                        
                        if prediction and probabilities is not None:
                            # Display prediction
                            col1, col2 = st.columns([3, 1])
                            
                            with col1:
                                if prediction == "Human":
                                    st.success(f"🎯 Prediction: **{prediction} Written**")
                                else:
                                    st.error(f"🎯 Prediction: **{prediction} Written**")
                            
                            with col2:
                                confidence = max(probabilities)
                                st.metric("Confidence", f"{confidence:.1%}")
                            
                            # Create probability chart
                            st.subheader("📊 Prediction Probabilities")
                            
                            # Detailed probabilities
                            col1, col2 = st.columns(2)
                            with col1:
                                st.metric("AI", f"{probabilities[1]:.1%}")
                            with col2:
                                st.metric("Human", f"{probabilities[0]:.1%}")
                            
                            # Bar chart
                            class_names = ['Human', 'AI']
                            prob_df = pd.DataFrame({
                                'Written by': class_names,
                                'Probability': probabilities
                            })
                            st.bar_chart(prob_df.set_index('Written by'), height=300)
                            
                        else:
                            st.error("Failed to make prediction")
                else:
                    st.warning("Please enter some text to classify!")
        else:
            st.error("No models available for prediction.")
    else:
        st.warning("Models not loaded. Please check the model files.")

# ============================================================================
# BATCH PROCESSING PAGE
# ============================================================================

elif page == "📁 Batch Processing":
    st.header("📁 Upload File for Batch Processing")
    st.markdown("Upload a CSV, PDF, word, or text file to process multiple texts at once.")

    def extract_text(file):
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    
    if models:
        available_models = get_available_models(models)
        
        if available_models:
            # File upload
            uploaded_file = st.file_uploader(
                "Choose a file",
                type=['txt', 'csv', 'pdf', 'docx'],
                help="Upload a .txt file (one text per line), .csv file (text in first column), .pfd file, or .docx file (one text per paragraph)"
            )
            
            if uploaded_file:
                # Model selection
                model_choice = st.selectbox(
                    "Choose model for batch processing:",
                    options=[model[0] for model in available_models],
                    format_func=lambda x: next(model[1] for model in available_models if model[0] == x)
                )
                
                # Process file
                if st.button("📊 Process File"):
                    try:
                        file_extension = pathlib.Path(uploaded_file.name).suffix
                        # Read file content
                        if file_extension == ".txt":
                            content = str(uploaded_file.read(), "utf-8")
                            texts = [line.strip() for line in content.split('\n') if line.strip()]
                        elif file_extension == ".csv":
                            df = pd.read_csv(uploaded_file)
                            texts = df.iloc[:, 0].astype(str).tolist()
                        elif file_extension == ".pdf":
                            pdf = pdfplumber.open(uploaded_file)
                            texts = ''
                            for page in pdf.pages:
                                texts += page.extract_text()
                            texts = texts.split("\t")
                        else: # docx
                            texts = extract_text(uploaded_file)
                            texts = [line.strip() for line in texts.split('\n') if line.strip()]
                        
                        if not texts:
                            st.error("No text found in file")
                        else:
                            st.info(f"Processing {len(texts)} texts...")
                            
                            # Process all texts
                            results = []
                            progress_bar = st.progress(0)
                            
                            for i, text in enumerate(texts):
                                if text.strip():
                                    prediction, probabilities = make_prediction(text, model_choice, models)
                                    
                                    if prediction and probabilities is not None:
                                        results.append({
                                            'Text': text[:100] + "..." if len(text) > 100 else text,
                                            'Full_Text': text,
                                            'Prediction': prediction,
                                            'Confidence': f"{max(probabilities):.1%}",
                                            'Human_Prob': f"{probabilities[0]:.1%}",
                                            'AI_Prob': f"{probabilities[1]:.1%}"
                                        })
                                
                                progress_bar.progress((i + 1) / len(texts))
                            
                            if results:
                                # Display results
                                st.success(f"✅ Processed {len(results)} texts successfully!")
                                
                                results_df = pd.DataFrame(results)
                                
                                # Summary statistics
                                st.subheader("📊 Summary Statistics")
                                col1, col2, col3, col4 = st.columns(4)
                                
                                positive_count = sum(1 for r in results if r['Prediction'] == 'AI')
                                negative_count = len(results) - positive_count
                                avg_confidence = np.mean([float(r['Confidence'].strip('%')) for r in results])
                                
                                with col1:
                                    st.metric("Total Processed", len(results))
                                with col2:
                                    st.metric("AI", positive_count)
                                with col3:
                                    st.metric("Human", negative_count)
                                with col4:
                                    st.metric("Avg Confidence", f"{avg_confidence:.1f}%")
                                
                                # Results preview
                                st.subheader("📋 Results Preview")
                                st.dataframe(
                                    results_df[['Text', 'Prediction', 'Confidence']],
                                    use_container_width=True
                                )
                                
                                # Download option
                                csv = results_df.to_csv(index=False)
                                st.download_button(
                                    label="📥 Download Full Results",
                                    data=csv,
                                    file_name=f"predictions_{model_choice}_{uploaded_file.name}.csv",
                                    mime="text/csv"
                                )
                            else:
                                st.error("No valid texts could be processed")
                                
                    except Exception as e:
                        st.error(f"Error processing file: {e}")
            else:
                st.info("Please upload a file to get started.")
                
                # Show example file formats
                with st.expander("📄 Example File Formats"):
                    st.markdown("""
                    **Text File (.txt):**
                    ```
                    This product is amazing!
                    Terrible quality, very disappointed
                    Great service and fast delivery
                    ```
                    
                    **CSV File (.csv):**
                    ```
                    text,category
                    "Amazing product, love it!",review
                    "Poor quality, not satisfied",review
                    ```
                    """)
        else:
            st.error("No models available for batch processing.")
    else:
        st.warning("Models not loaded. Please check the model files.")

# ============================================================================
# MODEL COMPARISON PAGE
# ============================================================================

elif page == "⚖️ Model Comparison":
    st.header("⚖️ Compare Models")
    st.markdown("Compare predictions from different models on the same text.")
    
    if models:
        available_models = get_available_models(models)
        
        if len(available_models) >= 2:
            # Text input for comparison
            comparison_text = st.text_area(
                "Enter text to compare models:",
                placeholder="Enter text to see how different models perform...",
                height=100
            )
            
            if st.button("📊 Compare All Models") and comparison_text.strip():
                st.subheader("🔍 Model Comparison Results")
                
                # Get predictions from all available models
                comparison_results = []
                
                for model_key, model_name in available_models:
                    prediction, probabilities = make_prediction(comparison_text, model_key, models)
                    
                    if prediction and probabilities is not None:
                        comparison_results.append({
                            'Model': model_name,
                            'Prediction': prediction,
                            'Confidence': f"{max(probabilities):.1%}",
                            'Human %': f"{probabilities[0]:.1%}",
                            'AI %': f"{probabilities[1]:.1%}",
                            'Raw_Probs': probabilities
                        })
                
                if comparison_results:
                    # Comparison table
                    comparison_df = pd.DataFrame(comparison_results)
                    st.table(comparison_df[['Model', 'Prediction', 'Confidence', 'Human %', 'AI %']])
                    
                    # Agreement analysis
                    predictions = [r['Prediction'] for r in comparison_results]
                    if len(set(predictions)) == 1:
                        st.success(f"✅ All models agree: **{predictions[0]} Written**")
                    else:
                        st.warning("⚠️ Models disagree on prediction")
                        for result in comparison_results:
                            model_name = result['Model'].split(' ')[1] if ' ' in result['Model'] else result['Model']
                            st.write(f"- {model_name}: {result['Prediction']}")
                    
                    # Side-by-side probability charts
                    st.subheader("📊 Detailed Probability Comparison")
                    
                    cols = st.columns(len(comparison_results))
                    
                    for i, result in enumerate(comparison_results):
                        with cols[i]:
                            model_name = result['Model']
                            st.write(f"**{model_name}**")
                            
                            chart_data = pd.DataFrame({
                                'Written by': ['Human', 'AI'],
                                'Probability': result['Raw_Probs']
                            })
                            st.bar_chart(chart_data.set_index('Written by'))
                    
                else:
                    st.error("Failed to get predictions from models")
        
        elif len(available_models) == 1:
            st.info("Only one model available. Use Single Prediction page for detailed analysis.")
            
        else:
            st.error("No models available for comparison.")
    else:
        st.warning("Models not loaded. Please check the model files.")

# ============================================================================
# MODEL INFO PAGE
# ============================================================================

elif page == "📊 Model Info":
    st.header("📊 Model Information")
    
    if models:
        st.success("✅ Models are loaded and ready!")
        
        # Model details
        st.subheader("🔧 Available Models")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            ### 📈 SVM
            **Type:** Support Vector Classification Model
            **Algorithm:** SVC
            **Features:** TF-IDF vectors (unigrams + bigrams)
            """)
            
        with col2:
            st.markdown("""
            ### 🎯 Decision Tree
            **Type:** Decision Tree Classification Model
            **Algorithm:** Decision Tree
            **Features:** TF-IDF vectors (unigrams + bigrams)
            """)
        
        with col3:
            st.markdown("""
            ### AdaBoost
            **Type:** Decision Tree Classification Model
            **Algorithm:** Decision Tree
            **Features:** TF-IDF vectors (unigrams + bigrams)
            """)
        
        # Feature engineering info
        st.subheader("🔤 Feature Engineering")
        st.markdown("""
        **Vectorization:** TF-IDF (Term Frequency-Inverse Document Frequency)
        - **Max Features:** 5,000 most important terms
        - **N-grams:** Unigrams (1-word) and Bigrams (2-word phrases)
        - **Stop Words:** English stop words removed
        """)
        
        # File status
        st.subheader("📁 Model Files Status")
        file_status = []
        
        files_to_check = [
            ("ai_detection_pipeline.pkl", "Complete SVM Pipeline", models.get('pipeline_available', False)),
            ("tfidf_vectorizer.pkl", "TF-IDF Vectorizer", models.get('vectorizer_available', False)),
            ("svm_model.pkl", "SVM Classifier", models.get('svm_available', False)),
            ("dt_model.pkl", "Decision Tree Classifier", models.get('dt_available', False)),
            ("ada_model.pkl", "AdaBoost Classifier", models.get('ada_available', False))
        ]
        
        for filename, description, status in files_to_check:
            file_status.append({
                "File": filename,
                "Description": description,
                "Status": "✅ Loaded" if status else "❌ Not Found"
            })
        
        st.table(pd.DataFrame(file_status))
        
        # Training information
        st.subheader("📚 Training Information")
        st.markdown("""
        **Dataset:** Product Review AI Detection Analysis
        - **Classes:** Human and AI written
        - **Preprocessing:** Text cleaning, tokenization, TF-IDF vectorization
        - **Training:** All models trained on same feature set for fair comparison
        """)
        
    else:
        st.warning("Models not loaded. Please check model files in the 'models/' directory.")

# ============================================================================
# HELP PAGE
# ============================================================================

elif page == "❓ Help":
    st.header("❓ How to Use This App")
    
    with st.expander("🔮 Single Prediction"):
        st.write("""
        1. **Select a model** from the dropdown (SVM, Decision Tree, or AdaBoost)
        2. **Enter text** in the text area
        3. **Click 'Predict'** to get AI detection results
        4. **View results:** prediction, confidence score, and probability breakdown
        5. **Try examples:** Use the provided example texts to test the models
        """)
    
    with st.expander("📁 Batch Processing"):
        st.write("""
        1. **Prepare your file:**
           - **.txt file:** One text per line
           - **.csv file:** Text in the first column
        2. **Upload the file** using the file uploader
        3. **Select a model** for processing
        4. **Click 'Process File'** to analyze all texts
        5. **Download results** as CSV file with predictions and probabilities
        """)
    
    with st.expander("⚖️ Model Comparison"):
        st.write("""
        1. **Enter text** you want to analyze
        2. **Click 'Compare All Models'** to get predictions from both models
        3. **View comparison table** showing predictions and confidence scores
        4. **Analyze agreement:** See if models agree or disagree
        5. **Compare probabilities:** Side-by-side probability charts
        """)
    
    with st.expander("🔧 Troubleshooting"):
        st.write("""
        **Common Issues and Solutions:**
        
        **Models not loading:**
        - Ensure model files (.pkl) are in the 'models/' directory
        - Check that required files exist:
          - tfidf_vectorizer.pkl (required)
          - ai_detection_pipeline.pkl (for SVM pipeline)
          - svm_model.pkl (for SVM individual)
          - dt_model.pkl (for Decision Tree model)
          - ada_model.pkl (for AdaBoost model)
        
        **Prediction errors:**
        - Make sure input text is not empty
        - Try shorter texts if getting memory errors
        - Check that text contains readable characters
        
        **File upload issues:**
        - Ensure file format is .txt or .csv
        - Check file encoding (should be UTF-8)
        - Verify CSV has text in the first column
        """)
    
    # System information
    st.subheader("💻 Your Project Structure")
    st.code("""
    streamlit_ml_app/
    ├── app.py                              # Main application
    ├── requirements.txt                    # Dependencies
    ├── models/                            # Model files
    │   ├── ai_detection_pipeline.pkl      # SVM complete pipeline
    │   ├── tfidf_vectorizer.pkl           # Feature extraction
    │   ├── svm_model.pkl                  # SVM classifier
    │   └── dt_model.pkl                   # Decision Tree classifier    
    │   └── ada_model.pkl                  # AdaBoost classifier   
    └── sample_data/                       # Sample files
        ├── sample_texts.txt
        └── sample_data.csv
    """)

# ============================================================================
# FOOTER
# ============================================================================

st.sidebar.markdown("---")
st.sidebar.markdown("### 📚 App Information")
st.sidebar.info("""
**ML Text Classification App**
Built with Streamlit

**Models:** 
- 📈 SVM
- 🎯 Decision Tree
- 🎯 AdaBoost

**Framework:** scikit-learn
**Deployment:** Streamlit Cloud Ready
""")

st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666666;'>
    Built with ❤️ using Streamlit | Machine Learning Text AI Detection Demo | By Emily Novak<br>
    <small>Project 1 of the course **Introduction to Large Language Models**</small><br>
    <small>This app predicts whether text is written by an AI or human using trained ML models</small>
</div>
""", unsafe_allow_html=True)